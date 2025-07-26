"""
Synapse Data Platform - Documentation Generator

Automated documentation generation for the Synapse project.
"""

from docx import Document
from docx.shared import Pt

# Create document
doc = Document()

doc.add_heading('Synapse Data Platform Documentation', 0)

doc.add_heading('1. Project Overview', level=1)
doc.add_paragraph(
    'This web application allows users to upload Excel or CSV files, analyze them, and display an interactive summary of the data. '
    'It includes file upload functionality, data processing with pandas, interactive visualization, and missing value handling.'
)

doc.add_heading('2. Technical Foundation', level=1)
doc.add_heading('A. Django Framework', level=2)
doc.add_paragraph(
    'Django is a high-level web framework for developing Python web applications. '
    'It uses the MTV (Model-Template-View) pattern. In this project, the main logic is in the home(request) view, '
    'which handles file upload and data analysis.'
)
doc.add_heading('B. Data Processing with pandas', level=2)
doc.add_paragraph(
    'pandas is the standard library for tabular data manipulation and analysis in Python. '
    'It is used to read files, calculate statistics, detect missing values, and generate summaries.'
)
doc.add_heading('C. Interactive Visualization', level=2)
doc.add_paragraph(
    'Visualization is performed on the web using JavaScript and canvas, allowing users to select columns and chart types.'
)
doc.add_heading('D. File Management', level=2)
doc.add_paragraph(
    'The HTML form allows file uploads. Django receives and temporarily stores files for analysis.'
)

doc.add_heading('3. Application Flow', level=1)
doc.add_paragraph(
    '1. User uploads a file through the web interface.\n'
    '2. Backend processes the file with pandas.\n'
    '3. Statistics, missing values, and visualizations are displayed.\n'
    '4. User can remove rows with NaN values and re-analyze.'
)

doc.add_heading('4. Code Documentation', level=1)
doc.add_paragraph('Below is an example of the main view with documentation:')

code_example = '''
# Main view that handles file upload, analysis, and data cleaning
from django.shortcuts import render
from django.core.files.storage import FileSystemStorage
import pandas as pd
import numpy as np
import os
from django.conf import settings

def home(request):
    # If user requests to remove rows with NaN (via hidden form)
    if request.method == 'POST' and request.POST.get('action') == 'remove_nan':
        filename = request.POST.get('filename')
        file_path = os.path.join(settings.MEDIA_ROOT, filename)
        ...
'''

p = doc.add_paragraph()
p.add_run(code_example).font.name = 'Courier New'
p.add_run('\n... (see complete code in views.py)')

# Save document
output_filename = 'Synapse_Documentation.docx'
doc.save(output_filename)
print(f'Documentation generated: {output_filename}') 